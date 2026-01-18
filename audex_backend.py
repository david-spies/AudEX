# app.py - Main Flask Application
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import whisper
from pathlib import Path
import tempfile
from docx import Document
from fpdf import FPDF
import uuid
from moviepy.editor import VideoFileClip
import logging

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Supported file extensions
AUDIO_EXTENSIONS = {'.mp3', '.wav', '.m4a', '.flac', '.ogg', '.webm'}
VIDEO_EXTENSIONS = {'.mp4', '.avi', '.mov', '.mkv', '.webm'}
ALLOWED_EXTENSIONS = AUDIO_EXTENSIONS | VIDEO_EXTENSIONS

# Load Whisper model (using base model for balance of speed/accuracy)
# Options: tiny, base, small, medium, large
logger.info("Loading Whisper model...")
model = whisper.load_model("base")
logger.info("Whisper model loaded successfully")


def allowed_file(filename):
    """Check if file extension is allowed"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def extract_audio_from_video(video_path, output_path):
    """Extract audio from video file"""
    try:
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(output_path, logger=None)
        video.close()
        return True
    except Exception as e:
        logger.error(f"Error extracting audio: {str(e)}")
        return False


def transcribe_audio(audio_path):
    """Transcribe audio file using Whisper"""
    try:
        logger.info(f"Transcribing audio file: {audio_path}")
        result = model.transcribe(audio_path, fp16=False)
        return {
            'success': True,
            'text': result['text'],
            'segments': result.get('segments', []),
            'language': result.get('language', 'unknown')
        }
    except Exception as e:
        logger.error(f"Transcription error: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def create_txt(text, output_path):
    """Create TXT file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def create_docx(text, output_path):
    """Create DOCX file"""
    doc = Document()
    doc.add_heading('Audio Transcription', 0)
    doc.add_paragraph(text)
    doc.save(output_path)


def create_pdf(text, output_path):
    """Create PDF file"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Audio Transcription", ln=True, align='C')
    pdf.ln(10)
    
    # Add text content
    pdf.set_font("Arial", size=12)
    # Split text into lines to fit PDF width
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, txt=line)
    
    pdf.output(output_path)


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'service': 'AudEX'}), 200


@app.route('/api/transcribe', methods=['POST'])
def transcribe():
    """Handle audio/video transcription"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported'}), 400
    
    try:
        # Generate unique filename
        file_id = str(uuid.uuid4())
        original_ext = Path(file.filename).suffix.lower()
        filename = secure_filename(f"{file_id}{original_ext}")
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save uploaded file
        file.save(filepath)
        logger.info(f"File saved: {filepath}")
        
        # If video, extract audio
        audio_path = filepath
        if original_ext in VIDEO_EXTENSIONS:
            audio_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.wav")
            logger.info("Extracting audio from video...")
            if not extract_audio_from_video(filepath, audio_path):
                return jsonify({'error': 'Failed to extract audio from video'}), 500
        
        # Transcribe audio
        result = transcribe_audio(audio_path)
        
        # Cleanup temporary files
        if os.path.exists(filepath):
            os.remove(filepath)
        if audio_path != filepath and os.path.exists(audio_path):
            os.remove(audio_path)
        
        if result['success']:
            return jsonify({
                'success': True,
                'transcription': result['text'],
                'language': result['language'],
                'segments': result.get('segments', [])
            }), 200
        else:
            return jsonify({'error': result['error']}), 500
            
    except Exception as e:
        logger.error(f"Transcription error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/export', methods=['POST'])
def export():
    """Export transcription to various formats"""
    data = request.get_json()
    
    if not data or 'text' not in data or 'format' not in data:
        return jsonify({'error': 'Missing required fields'}), 400
    
    text = data['text']
    format_type = data['format'].lower()
    
    if format_type not in ['txt', 'docx', 'pdf']:
        return jsonify({'error': 'Unsupported export format'}), 400
    
    try:
        # Generate unique filename
        file_id = str(uuid.uuid4())
        output_filename = f"transcription_{file_id}.{format_type}"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        # Create file based on format
        if format_type == 'txt':
            create_txt(text, output_path)
        elif format_type == 'docx':
            create_docx(text, output_path)
        elif format_type == 'pdf':
            create_pdf(text, output_path)
        
        # Send file
        return send_file(
            output_path,
            as_attachment=True,
            download_name=f"transcription.{format_type}"
        )
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/formats', methods=['GET'])
def get_supported_formats():
    """Return supported file formats"""
    return jsonify({
        'audio': list(AUDIO_EXTENSIONS),
        'video': list(VIDEO_EXTENSIONS),
        'export': ['txt', 'docx', 'pdf']
    }), 200


if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000, threaded=True)
